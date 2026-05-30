import io
from pathlib import Path
from typing import Dict, List

import dxpdf
import pandas as pd
from docx import Document
from loguru import logger

from app.models.generated_document import GeneratedDocumentCreate
from app.services.generated_document_service import GeneratedDocumentService


class DocumentGeneratorService:
    OUTPUT_DIR = Path.cwd() / "templates/output"
    OUTPUT_DIR.mkdir(exist_ok=True)
    PDF_DIR = Path.cwd() / "templates/pdf"

    def excel_to_dicts(self, excel_content: bytes) -> List[Dict]:
        """Excel bytes → list[dict]"""
        df = pd.read_excel(io.BytesIO(excel_content))
        return df.to_dict("records")

    def fill_docx_template(self, template_path: str, data: Dict, output_path: str):
        """Заполняет один DOCX"""
        doc = Document(template_path)

        # Параграфы
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{key}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

        # Таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder = f"{{{key}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

        doc.save(output_path)

    async def generate_documents(
        self,
        generated_doc_service: GeneratedDocumentService,
        template_path: str,
        data_list: List[Dict],
        process_id: str,
    ) -> List[str]:
        output_dir = self.OUTPUT_DIR / process_id
        output_dir.mkdir(exist_ok=True)

        generated_paths = []

        for i, data in enumerate(data_list):
            # Сначала создаём запись в БД (без file_path)
            doc_record = await generated_doc_service.create(
                GeneratedDocumentCreate(
                    gen_process_id=process_id,
                    file_path="",  # временно
                )
            )

            # Имя файла = id из БД
            filename = f"{doc_record.id}.docx"
            output_path = output_dir / filename

            # Генерируем документ
            self.fill_docx_template(template_path, data, str(output_path))

            # Обновляем запись с реальным путём
            doc_record.file_path = str(output_path)
            await generated_doc_service.update(doc_record.id, doc_record)

            generated_paths.append(str(output_path))
            logger.info(f"Создан документ: {output_path} (id={doc_record.id})")

        return generated_paths

    def convert_to_pdf_sync(self, input_path: str) -> str:
        """Синхронно конвертирует DOCX в PDF через dxpdf."""
        input_path_obj = Path(input_path).resolve()
        output_path = self.PDF_DIR / f"{input_path_obj.stem}.pdf"

        # Убедимся, что папка для PDF существует
        self.PDF_DIR.mkdir(parents=True, exist_ok=True)

        # Конвертируем
        dxpdf.convert_file(str(input_path_obj), str(output_path))

        logger.info(f"PDF создан с помощью dxpdf: {output_path}")
        return str(output_path)
